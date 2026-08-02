"""
Demo: Review otomatis naskah .docx mahasiswa
Bukti bahwa python-docx bisa membaca format, struktur, dan konten file Word.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter, defaultdict
import re

DOCX_PATH = r"naskah PI\PI_Muhammad_Fauza_BAB1-3.docx"


def emu_to_cm(emu):
    """Convert EMU (English Metric Units) to cm."""
    if emu is None:
        return None
    return round(emu / 360000, 2)


def pt_to_float(pt_val):
    """Convert Pt object to float."""
    if pt_val is None:
        return None
    return round(pt_val.pt, 1)


def analyze_document(path: str) -> dict:
    doc = Document(path)
    results = {
        "margins": {},
        "fonts": Counter(),
        "font_sizes": Counter(),
        "line_spacings": Counter(),
        "headings": [],
        "paragraph_count": 0,
        "word_count": 0,
        "page_estimate": 0,
        "bab_found": [],
        "has_abstrak": False,
        "has_kata_pengantar": False,
        "has_daftar_isi": False,
        "has_daftar_pustaka": False,
        "tables_count": len(doc.tables),
        "images_count": 0,
        "alignment_stats": Counter(),
    }

    # ── 1. CEK MARGIN ──
    section = doc.sections[0]
    results["margins"] = {
        "atas": emu_to_cm(section.top_margin),
        "bawah": emu_to_cm(section.bottom_margin),
        "kiri": emu_to_cm(section.left_margin),
        "kanan": emu_to_cm(section.right_margin),
    }

    # ── 2. CEK FONT, UKURAN, SPASI, ALIGNMENT ──
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        results["paragraph_count"] += 1
        results["word_count"] += len(text.split())

        # Alignment
        if para.alignment is not None:
            results["alignment_stats"][str(para.alignment)] += 1

        # Line spacing
        if para.paragraph_format.line_spacing is not None:
            spacing = para.paragraph_format.line_spacing
            if isinstance(spacing, float) or isinstance(spacing, int):
                results["line_spacings"][f"{spacing}x"] += 1
            else:
                results["line_spacings"][f"{pt_to_float(spacing)}pt"] += 1

        # Font info dari setiap run
        for run in para.runs:
            if run.font.name:
                results["fonts"][run.font.name] += 1
            if run.font.size:
                results["font_sizes"][f"{pt_to_float(run.font.size)}pt"] += 1

        # ── 3. DETEKSI STRUKTUR (BAB, ABSTRAK, dll) ──
        text_lower = text.lower()

        # Deteksi BAB
        bab_match = re.match(r"bab\s+(i{1,3}|iv|v)\b", text_lower)
        if bab_match:
            results["bab_found"].append(text[:60])

        # Deteksi section penting
        if "abstrak" in text_lower and len(text) < 30:
            results["has_abstrak"] = True
        if "kata pengantar" in text_lower and len(text) < 30:
            results["has_kata_pengantar"] = True
        if "daftar isi" in text_lower and len(text) < 30:
            results["has_daftar_isi"] = True
        if "daftar pustaka" in text_lower and len(text) < 30:
            results["has_daftar_pustaka"] = True

        # Deteksi heading styles
        if para.style and "heading" in para.style.name.lower():
            results["headings"].append({
                "level": para.style.name,
                "text": text[:80],
            })

    # Hitung gambar (inline shapes)
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            results["images_count"] += 1

    # Estimasi halaman (rough: ~250 kata per halaman)
    results["page_estimate"] = max(1, results["word_count"] // 250)

    return results


def print_review(results: dict):
    print()
    print("=" * 65)
    print("📋 LAPORAN REVIEW OTOMATIS NASKAH PI")
    print("   Powered by python-docx")
    print("=" * 65)

    # ── MARGIN ──
    print("\n📏 1. MARGIN HALAMAN")
    print("─" * 40)
    expected_margins = {"atas": 3.0, "bawah": 3.0, "kiri": 4.0, "kanan": 3.0}
    for pos, val in results["margins"].items():
        expected = expected_margins.get(pos, "?")
        status = "✅" if val and abs(val - expected) < 0.2 else "❌"
        print(f"  {status} Margin {pos}: {val} cm (standar: {expected} cm)")

    # ── FONT ──
    print("\n🔤 2. FONT YANG DIGUNAKAN")
    print("─" * 40)
    for font, count in results["fonts"].most_common(5):
        is_tnr = "times new roman" in font.lower()
        status = "✅" if is_tnr else "⚠️"
        print(f"  {status} {font}: {count} penggunaan")

    # ── UKURAN FONT ──
    print("\n📐 3. UKURAN FONT")
    print("─" * 40)
    for size, count in results["font_sizes"].most_common(5):
        is_12 = "12.0" in size
        status = "✅" if is_12 else "ℹ️"
        print(f"  {status} {size}: {count} penggunaan")

    # ── SPASI ──
    print("\n📏 4. SPASI BARIS")
    print("─" * 40)
    if results["line_spacings"]:
        for spacing, count in results["line_spacings"].most_common(5):
            print(f"  ℹ️ {spacing}: {count} paragraf")
    else:
        print("  ℹ️ Menggunakan spasi default dokumen")

    # ── STATISTIK ──
    print("\n📊 5. STATISTIK DOKUMEN")
    print("─" * 40)
    print(f"  📄 Estimasi halaman: ~{results['page_estimate']} halaman", end="")
    print(f" {'✅ (≥40)' if results['page_estimate'] >= 40 else '⚠️ (target: min 40)'}")
    print(f"  📝 Total paragraf: {results['paragraph_count']}")
    print(f"  📖 Total kata: {results['word_count']:,}")
    print(f"  📊 Jumlah tabel: {results['tables_count']}")
    print(f"  🖼️ Jumlah gambar: {results['images_count']}")

    # ── KELENGKAPAN STRUKTUR ──
    print("\n📑 6. KELENGKAPAN STRUKTUR")
    print("─" * 40)

    checks = [
        ("Abstrak", results["has_abstrak"]),
        ("Kata Pengantar", results["has_kata_pengantar"]),
        ("Daftar Isi", results["has_daftar_isi"]),
        ("Daftar Pustaka", results["has_daftar_pustaka"]),
    ]
    for name, found in checks:
        print(f"  {'✅' if found else '❌'} {name}: {'Ditemukan' if found else 'Tidak ditemukan'}")

    print(f"\n  BAB yang terdeteksi ({len(results['bab_found'])}):")
    for bab in results["bab_found"]:
        print(f"    ✅ {bab}")

    # Cek BAB yang hilang
    expected_babs = ["bab i", "bab ii", "bab iii", "bab iv", "bab v"]
    found_lower = [b.lower() for b in results["bab_found"]]
    for bab in expected_babs:
        if not any(bab in f for f in found_lower):
            print(f"    ❌ {bab.upper()} tidak ditemukan")

    print("\n" + "=" * 65)
    print("✨ Review selesai!")
    print("=" * 65)


if __name__ == "__main__":
    print(f"📂 Membaca file: {DOCX_PATH}")
    results = analyze_document(DOCX_PATH)
    print_review(results)
